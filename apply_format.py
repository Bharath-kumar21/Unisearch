"""
Generate Project_Report.docx from the existing '4-2 documentation.docx' content,
restructured into the FORMAT.txt headings with Times New Roman font
(16pt headings, 14pt body).
"""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading_para(doc, text, level=1):
    """Add a heading in Times New Roman 16pt Bold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text):
    """Add body text in Times New Roman 14pt with 1.5 line spacing."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    from docx.shared import Pt as _Pt
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(doc, text):
    """Add a bullet point in Times New Roman 14pt."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.5
    return p

def page_break(doc):
    doc.add_page_break()

# ── build document ───────────────────────────────────────────────────────────

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

# Set narrow margins so more content fits per page
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ═══════════════════════ PAGE 1-2: 1. Introduction ═══════════════════════════

add_heading_para(doc, '1. Introduction')

add_body(doc,
    'Choosing the right university is one of the most important decisions in a student\'s academic journey. '
    'For many students, especially those applying through competitive entrance examinations, this decision is '
    'a complex evaluation process involving rank eligibility, tuition fees, placement opportunities, accreditation, '
    'branch availability, location preferences, and long-term career outcomes. In countries like India, where '
    'admissions are often driven by national and state-level entrance examinations, the challenge becomes even '
    'greater because of the sheer volume of institutions and the dynamic nature of cutoff trends.')

add_body(doc,
    'Most students today rely on a combination of unofficial websites, coaching institute lists, social media '
    'discussions, and peer recommendations to make their choices. While these sources can be helpful, they are '
    'rarely structured in a way that supports objective and data-driven decision-making. Information is often '
    'outdated, inconsistent across platforms, or incomplete for specific branches and categories.')

add_body(doc,
    'To address this gap, the "Integrated Information Platform for Indian Universities" (UniSearch India) was '
    'designed as a full-stack web application that simplifies the university selection process by combining '
    'verified data, intelligent filtering, and interactive analytics into a single unified interface. The platform '
    'enables students to input their exam type and rank, instantly discover eligible universities, and compare '
    'institutions across multiple criteria such as placements, fees, accreditation, and branch-specific cutoffs.')

add_body(doc,
    'At its core, the platform aims to answer a simple but critical question: "Given my rank, budget, and '
    'preferences, which universities are the best options for me?" This project is not just a directory of '
    'universities; it is a decision-support system. It integrates backend APIs, secure authentication, a '
    'structured database, and a modern React frontend to create a seamless user experience from initial search '
    'to final shortlist.')

add_body(doc,
    'The system architecture is designed to be modular and scalable, making it suitable not only as an academic '
    'project but also as a production-ready foundation that can be extended in the future with recommendation '
    'algorithms, predictive analytics, and personalized counseling features.')


# ═══════════════════════ PAGE 3: 1.1 Problem Statement ═══════════════════════

page_break(doc)

add_heading_para(doc, '1.1 Problem Statement')

add_body(doc,
    'The process of selecting a suitable university after entrance examinations is a major challenge for students '
    'due to the lack of a centralized, reliable, and personalized decision-making system. Students are required '
    'to evaluate multiple factors such as entrance rank eligibility, branch-wise cutoff trends, tuition fees, '
    'placement opportunities, accreditation status, and geographic preferences. However, this information is '
    'usually scattered across different sources including official university websites, counseling PDFs, coaching '
    'institute documents, and third-party portals.')

add_body(doc,
    'Existing systems and resources have several limitations. Most platforms provide generic lists of top colleges '
    'without considering the student\'s individual rank, preferred exam type, budget, or branch interest. '
    'Official websites may contain authentic data, but they are difficult to compare because each institution '
    'presents information in a different format. Coaching materials and static cutoff documents quickly become '
    'outdated and do not support dynamic filtering.')

add_body(doc,
    'Another key issue is the absence of an integrated comparison mechanism. Students often manually open '
    'multiple tabs and try to compare colleges by memory or handwritten notes, which is inefficient and '
    'error-prone. There is no standard system that allows them to compare universities side by side based on '
    'measurable indicators such as NAAC grade, placement percentage, annual fees, and branch-specific cutoff ranks.')

add_body(doc,
    'Therefore, the problem this project aims to solve is: how to design and implement an integrated University '
    'Recommendation and Comparison Platform that reduces information fragmentation, supports personalized '
    'rank-based filtering, and enables students to make informed, data-driven admission decisions efficiently '
    'and confidently.')


# ═══════════════════════ PAGE 4: 1.2 Objective of the project ════════════════

page_break(doc)

add_heading_para(doc, '1.2 Objective of the Project')

add_body(doc,
    'The primary objective of this project is to design and develop a University Recommendation and Comparison '
    'Platform that helps students make informed admission decisions based on their entrance exam rank, preferences, '
    'and key institutional parameters. The system is intended to solve the common problem of information '
    'fragmentation and decision uncertainty by providing a centralized, data-driven, and user-friendly web '
    'application.')

add_body(doc, 'The specific objectives are:')

objectives = [
    'To build a rank-based recommendation mechanism that allows students to input their exam type and rank, and then receive a list of universities that are realistically attainable.',
    'To implement a comprehensive filtering and search system enabling users to refine results based on state, fee range, accreditation grade, institution type, and placement percentage.',
    'To provide a detailed university profile module where users can access branch-wise cutoff ranks, placement statistics, annual fees, and accreditation details.',
    'To develop a side-by-side comparison feature that allows students to compare up to three universities simultaneously across important metrics.',
    'To include an analytics dashboard with visual charts providing high-level insights such as state-wise distribution, NAAC grade spread, and fee-vs-placement correlations.',
    'To implement secure and scalable full-stack architecture using React, Express, and MongoDB with JWT-based authentication.',
    'To ensure the platform is deployment-ready with a decoupled architecture allowing independent hosting of frontend and backend.',
]
for obj in objectives:
    add_bullet(doc, obj)


# ═══════════════════════ PAGE 5: 1.3 Purpose of the project ══════════════════

page_break(doc)

add_heading_para(doc, '1.3 Purpose of the Project')

add_body(doc,
    'The purpose of this project is to develop a centralized and intelligent University Recommendation and '
    'Comparison Platform that simplifies the complex process of selecting a suitable university after entrance '
    'examinations. Every year, students face significant difficulty in deciding where to apply because admission '
    'decisions depend on multiple factors such as rank eligibility, branch-wise cutoffs, tuition fees, placement '
    'opportunities, accreditation, and location preferences.')

add_body(doc,
    'A major purpose of the system is to reduce the burden of manual research. Traditionally, students and '
    'parents spend hours collecting data from university websites, counseling PDFs, and educational portals, then '
    'try to compare institutions manually. The project aims to replace this inefficient method with an automated '
    'and interactive solution that allows users to enter their exam rank and instantly discover relevant universities.')

add_body(doc,
    'Another important purpose is to promote data-driven decision-making. Many students choose colleges based '
    'on reputation or peer suggestions without analyzing measurable factors. This platform encourages informed '
    'choices by presenting detailed university profiles including branch-specific cutoff ranks, fees, placements, '
    'and accreditation status. The analytics dashboard offers visual insights into the broader education landscape.')

add_body(doc,
    'The project also serves the purpose of increasing accessibility and fairness in educational guidance. Not '
    'all students have access to professional counseling services. By making reliable and organized information '
    'available through a free and user-friendly web platform, the project supports equal access to guidance and '
    'empowers students to make confident decisions independently.')


# ═══════════════════════ PAGE 6: 1.4 Scope of the project ════════════════════

page_break(doc)

add_heading_para(doc, '1.4 Scope of the Project')

add_body(doc,
    'The scope of this project covers the design and development of a full-stack University Recommendation '
    'and Comparison Platform that supports students in selecting suitable universities based on their entrance '
    'exam rank, preferences, and institutional data. The system provides an end-to-end solution for university '
    'discovery, filtering, comparison, and shortlisting.')

add_body(doc,
    'At the data level, the project includes collecting and maintaining structured university information such '
    'as institution name, location, accreditation, fees, placement percentage, and branch-wise cutoff ranks. '
    'The dataset is initialized from JSON and spreadsheet sources, and a seeding process populates the MongoDB '
    'database. Missing records are handled by applying fallback estimations.')

add_body(doc,
    'At the backend level, the scope includes building a RESTful API using Node.js and Express with an '
    'Authentication Module (signup/login with JWT), a User Module (protected profile endpoints), and a University '
    'Module (list all and fetch by ID). At the frontend level, the scope covers a React SPA with landing page, '
    'university listing, details page, comparison page, analytics dashboard, and authentication pages.')

add_body(doc,
    'Out of scope for the current version: category-wise reservation and quota-based cutoff prediction, '
    'real-time counseling round updates, AI/ML-based recommendation engine, admin panel for dynamic data '
    'management, and mobile application support.')


# ═══════════════════════ PAGE 6-7: 1.5 Benefits of the project ═══════════════

add_heading_para(doc, '1.5 Benefits of the Project')

add_body(doc,
    'The platform provides several important benefits for students, institutions, and the overall admission process:')

benefits = [
    'Simplified Decision-Making: Students can access all relevant university information in one place, saving time and reducing confusion during tight counseling deadlines.',
    'Personalized Recommendations: Unlike generic college lists, the platform tailors results based on the user\'s rank, exam type, and selected filters, ensuring students focus only on achievable options.',
    'Data-Driven Comparison: Detailed profiles and side-by-side comparison enable evaluation based on measurable factors such as placement percentage, fees, NAAC grade, and branch-wise cutoff ranks.',
    'Improved Accessibility: Free and structured access to essential admission data ensures all students, regardless of background, can explore opportunities independently.',
    'Visual Insights and Analytics: Dashboards with charts help users understand trends such as state distribution of universities, accreditation levels, and fee-placement relationships.',
    'Technical Scalability: The decoupled architecture using React, Express, and MongoDB allows independent deployment and future extension with features like AI recommendations.',
]
for b in benefits:
    add_bullet(doc, b)


# ═══════════════════════ PAGE 8-10: 2. Review Of Literature ══════════════════

page_break(doc)

add_heading_para(doc, '2. Review of Literature')

add_body(doc,
    'The process of university selection and admission guidance has been widely discussed in academic research '
    'and practical implementations. This section reviews existing systems, methodologies, and technological '
    'approaches related to college recommendation, decision support systems, and educational data platforms.')

add_body(doc,
    'Traditional Admission Guidance Systems: Historically, university selection has relied on manual counseling, '
    'career counselors, coaching institutes, and printed materials such as brochures and cutoff lists. Printed '
    'counseling documents provide valuable historical data but lack interactivity. Students must manually interpret '
    'these documents and compare multiple institutions. Access to professional counseling is often limited to '
    'urban or economically privileged students, creating an imbalance.')

add_body(doc,
    'Existing Online Educational Portals: Websites such as educational portals and ranking platforms offer lists '
    'of institutions, placement statistics, and user reviews. However, most provide static rankings or generalized '
    'lists that do not consider individual student profiles. Many portals rely on proprietary ranking algorithms '
    'that lack transparency. Some platforms restrict advanced features behind paid subscriptions, limiting accessibility.')

add_body(doc,
    'Decision Support Systems in Education: DSS platforms improve decision quality by providing structured data, '
    'analytical tools, and interactive interfaces. Many systems use multi-criteria decision-making (MCDM) '
    'techniques. However, many academic DSS models remain theoretical or limited in real-world deployment. The '
    'proposed project builds upon DSS principles by providing a practical, web-based implementation.')

add_body(doc,
    'Recommender Systems and Personalization: In educational applications, rank-based recommendation is a form '
    'of content-based filtering, where universities are matched based on cutoff criteria and user input. The '
    'proposed platform adopts a rule-based recommendation approach using rank and cutoff data to provide '
    'immediate suggestions without requiring complex machine learning models.')

add_body(doc,
    'Data Visualization in Educational Platforms: Research in HCI shows that visual representations improve '
    'comprehension, reduce cognitive load, and support better decision-making. The proposed project incorporates '
    'interactive dashboards using Recharts, enabling users to explore trends such as distribution of universities '
    'across states, accreditation levels, and relationships between fees and placement outcomes.')

add_body(doc,
    'Limitations of Existing Approaches: Despite advancements, several limitations remain — lack of '
    'personalization, data fragmentation, limited comparison features, poor accessibility, and static data '
    'representation. The proposed system addresses these gaps by providing rank-based personalized recommendations, '
    'advanced filtering, detailed profiles, side-by-side comparison, interactive dashboards, and secure authentication.')


# ═══════════════════════ PAGE 11: 3. SRA ═════════════════════════════════════

page_break(doc)

add_heading_para(doc, '3. Software Requirement Analysis')

add_body(doc,
    'The Software Requirement Analysis (SRA) phase focuses on understanding the problem domain, identifying '
    'user needs, defining system requirements, and establishing a clear foundation for design and implementation. '
    'This section is divided into two major parts: Domain Analysis and Requirement Analysis.')


# ═══════════════════════ PAGE 11-12: 3.1 Domain Analysis ═════════════════════

add_heading_para(doc, '3.1 Domain Analysis')

add_body(doc,
    'The project falls under the domain of Educational Technology (EdTech), specifically targeting university '
    'admission guidance systems. In India, university admissions are highly competitive and based on standardized '
    'entrance exams such as JEE, EAMCET, KCET, and others. Each university and branch has specific cutoff ranks '
    'that change every year.')

add_body(doc, 'Stakeholders in the Domain:')
stakeholders = [
    'Primary Users (Students): Enter exam rank and preferences, search and filter universities, compare institutions, save shortlisted colleges.',
    'Secondary Users (Parents/Guardians): Assist in decision-making, analyze financial aspects such as fees and ROI.',
    'System Administrators (Future Scope): Manage university data, update cutoff ranks and metrics.',
    'Educational Institutions (Indirect): Benefit from visibility on the platform.',
]
for s in stakeholders:
    add_bullet(doc, s)

add_body(doc, 'Existing Domain Challenges:')
challenges = [
    'Information Fragmentation: Data is spread across multiple platforms.',
    'Lack of Standardization: Different formats for cutoff, fees, placements.',
    'Dynamic Nature of Data: Cutoffs change yearly.',
    'Limited Personalization: No rank-based filtering in many systems.',
    'Manual Comparison: Students compare data manually.',
]
for c in challenges:
    add_bullet(doc, c)

add_body(doc,
    'The system handles structured and semi-structured data including university data (name, location, NAAC '
    'grade), academic data (branch-wise cutoff ranks), financial data (tuition fees), outcome data (placement '
    'percentages), and user data (login credentials, saved universities).')


# ═══════════════════════ PAGE 13-20: 3.2 Requirement Analysis ════════════════

page_break(doc)

add_heading_para(doc, '3.2 Requirement Analysis')

add_body(doc, 'Functional Requirements:')

add_body(doc,
    'User Registration and Authentication: Users should be able to sign up with email and password. The system '
    'hashes passwords securely using bcrypt. Users log in and receive JWT tokens. Session is maintained using '
    'localStorage.')

add_body(doc,
    'Rank-Based University Recommendation: Users input exam type and rank. The system filters universities '
    'based on cutoff criteria and displays only eligible universities.')

add_body(doc,
    'University Search and Filtering: Filter by state, fee range, NAAC grade, placement percentage, exam type, '
    'and institution type. Support multi-select filters and dynamic filtering based on user input.')

add_body(doc,
    'University Details View: Display full university information including branch-wise cutoff table, fees, '
    'placements, and official website link.')

add_body(doc,
    'Comparison Feature: Users can select up to 3 universities for side-by-side comparison on metrics including '
    'fees, placements, accreditation, and rank.')

add_body(doc,
    'Dashboard and Analytics: Display interactive charts — state distribution (bar chart), NAAC grades (pie '
    'chart), and fees vs placements (scatter chart).')

add_body(doc,
    'Bookmarking/Shortlisting: Users can save universities, stored in localStorage, and displayed in the '
    'profile page.')

add_body(doc, 'Non-Functional Requirements:')

nfr = [
    'Performance: Pages load within 2-3 seconds; API responses within 1-2 seconds; caching via Context API reduces redundant calls.',
    'Scalability: Supports increasing users without degradation; independent scaling of frontend and backend.',
    'Security: bcrypt password hashing, JWT authentication, protected routes, prevention of injection and XSS attacks.',
    'Usability: Simple, clean, and intuitive interface; responsive design across desktop, tablet, and mobile.',
    'Reliability: Minimal downtime, graceful error handling, fallback mechanisms for missing data.',
    'Maintainability: Modular and clean architecture, reusable components, proper documentation.',
    'Portability: Works on all modern browsers (Chrome, Edge, Firefox); supports deployment on Vercel, Render, and MongoDB Atlas.',
]
for n in nfr:
    add_bullet(doc, n)

add_body(doc, 'System Requirements:')

add_body(doc,
    'Hardware Requirements: Minimum 4GB RAM, standard PC/Laptop, and internet connectivity.')

add_body(doc,
    'Software Requirements: Frontend — React.js with Create React App and Vanilla CSS. '
    'Backend — Node.js with Express.js. Database — MongoDB / MongoDB Atlas. '
    'Tools — VS Code, Postman (API testing), Git & GitHub.')

add_body(doc, 'User Requirements:')
user_reqs = [
    'Easy-to-use interface with minimal learning curve.',
    'Fast and accurate search results based on rank.',
    'Secure login and session management.',
    'Mobile-friendly and responsive design.',
]
for u in user_reqs:
    add_bullet(doc, u)

add_body(doc, 'Constraints and Assumptions:')
add_bullet(doc, 'Limited real-time data updates; no category-based reservation logic in the current version.')
add_bullet(doc, 'Dependency on dataset accuracy; users are assumed to know their exam rank; internet access is required.')


# ═══════════════════════ PAGE 21-25: 4. System Analysis Design ═══════════════

page_break(doc)

add_heading_para(doc, '4. System Analysis and Design')

add_body(doc,
    'This chapter presents the system analysis and design of the University Recommendation and Comparison '
    'Platform. It covers the overall system architecture, database schema design, and UML diagrams that model '
    'the system behavior and structure.')


# ═══════════════════════ 4.1 System Design ═══════════════════════════════════

add_heading_para(doc, '4.1 System Design')

add_body(doc,
    'The system adopts a client-server architecture following the three-tier model: Presentation Layer (React '
    'SPA), Application Layer (Express REST API), and Data Layer (MongoDB). The frontend runs on port 3000 and '
    'communicates with the backend on port 5000 via HTTP fetch requests.')

add_body(doc,
    'Architecture Overview: The React frontend delegates state management to the Context API through three '
    'providers — AuthContext (JWT-based session management), UniversityContext (API data fetching and caching), '
    'and SavedContext (localStorage-based bookmarking). The Express server handles route protections and interacts '
    'asynchronously with MongoDB using Mongoose ODM.')

add_body(doc, 'Backend API Design:')
api_routes = [
    'POST /api/auth/signup — Register new user with name, email, password. Returns JWT token and user object.',
    'POST /api/auth/login — Authenticate user with email and password. Returns JWT token and user object.',
    'GET /api/users/profile — Protected endpoint. Returns user profile (excluding password) using JWT verification.',
    'GET /api/universities — Returns all universities sorted by NIRF ranking (ascending).',
    'GET /api/universities/:id — Returns a single university by MongoDB ObjectId.',
]
for a in api_routes:
    add_bullet(doc, a)

add_body(doc,
    'Authentication Flow: The user signs up or logs in and receives a JWT token (24-hour expiry). The token is '
    'stored in localStorage. On page load, AuthContext reads the token and fetches /api/users/profile to restore '
    'the session. Protected routes redirect to /login if no valid session exists. Logout clears the token and '
    'resets state.')

add_body(doc, 'Database Schema Design:')

add_body(doc,
    'University Schema: Fields include name (String, required), location (String, required), ranking (Number, '
    'required), fees (String), average_placement (String), placement_percentage (String), website (String), '
    'established (Number), type (String — "Public" or "Private"), required_exam (String), '
    'average_rank_required (String), and branches (Array of objects with name and cutoffRank).')

add_body(doc,
    'User Schema: Fields include name (String, required, trimmed), email (String, required, unique, lowercase), '
    'password (String, required, bcrypt-hashed), joined (String, auto-generated), and reviews (Number, '
    'default: 0). Timestamps are enabled.')

add_body(doc, 'Frontend Component Architecture:')
components = [
    'Pages: Home.js (landing page with exam/rank search), Universities.js (filterable paginated listing), UniversityDetails.js (detailed view), Compare.js (side-by-side comparison), Dashboard.js (analytics charts), Login.js, SignUp.js, Profile.js.',
    'Components: Navbar.js (global navigation with search), Footer.js (site footer), UniversityCard.js (listing card with save toggle), SearchBar.js.',
    'Context Providers: AuthContext.js (authentication state), UniversityContext.js (data fetching/caching), SavedContext.js (bookmarked universities).',
]
for comp in components:
    add_bullet(doc, comp)


# ═══════════════════════ PAGE 26-27: 4.2 ER Diagram ══════════════════════════

page_break(doc)

add_heading_para(doc, '4.2 ER Diagram')

add_body(doc,
    'The Entity-Relationship diagram delineates two primary entities in the system: User and University. '
    'The User entity stores authentication credentials and profile information. The University entity stores '
    'all institutional data including academic metrics and branch-wise cutoff information.')

add_body(doc,
    'A conceptual one-to-many relationship exists where a User can save (bookmark) multiple Universities. '
    'In the current implementation, this relationship is managed client-side through localStorage via the '
    'SavedContext, storing an array of university IDs associated with the user session.')

add_body(doc, 'User Entity Attributes:')
add_bullet(doc, '_id (Primary Key, auto-generated), name, email (unique), password (hashed), joined, reviews, timestamps')

add_body(doc, 'University Entity Attributes:')
add_bullet(doc, '_id (Primary Key, auto-generated), name, location, ranking, fees, average_placement, placement_percentage, website, established, type, required_exam, average_rank_required, branches (embedded array)')

add_body(doc,
    'The Branch sub-entity is embedded within the University document as an array of objects, each containing '
    'a branch name (e.g., Computer Science, Electronics) and the corresponding cutoffRank. This denormalized '
    'design avoids joins and improves query performance in MongoDB.')

add_body(doc, '[ER Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 28: 4.3 UML Diagrams ═══════════════════════════

page_break(doc)

add_heading_para(doc, '4.3 UML Diagrams')

add_body(doc,
    'Unified Modeling Language (UML) diagrams are utilized to visualize the architecture of the UniSearch India '
    'platform. They blueprint both the static structural framework and the dynamic behavioral interactions within '
    'the application. The following subsections present nine types of UML diagrams that collectively '
    'describe the system\'s design from multiple perspectives.')


# ═══════════════════════ PAGE 29-30: 4.3.1 Class Diagram ═════════════════════

add_heading_para(doc, '4.3.1 Class Diagram')

add_body(doc,
    'The Class diagram highlights the core application components and their relationships. The primary classes '
    'are the University model (with attributes: name, location, ranking, fees, average_placement, '
    'placement_percentage, website, established, type, required_exam, average_rank_required; and an aggregation '
    'of Branch objects), the User model (with attributes: name, email, password, joined, reviews), and the '
    'Context classes (AuthContext with methods login(), signup(), logout(); UniversityContext with fetchAll(); '
    'SavedContext with save(), unsave()).')

add_body(doc,
    'UI component classes include Home (with searchByRank()), Universities (with applyFilters(), paginate()), '
    'UniversityDetails (with fetchById()), Compare (with addToCompare(), removeFromCompare()), and Dashboard '
    '(with renderCharts()). The Navbar component includes a global search method, and UniversityCard includes '
    'a toggleSave() method.')

add_body(doc, '[Class Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 31-32: 4.3.2 Use-case Diagram ══════════════════

page_break(doc)

add_heading_para(doc, '4.3.2 Use-case Diagram')

add_body(doc,
    'The Use-case diagram captures the primary actors and their interactions with the system. The main actors '
    'are the Student (User) and the System (Server + Database).')

add_body(doc, 'Actor: Student (Unauthenticated)')
add_bullet(doc, 'Search Universities by Rank and Exam Type')
add_bullet(doc, 'Browse and Filter University Listings')
add_bullet(doc, 'View University Details')
add_bullet(doc, 'Compare Universities (up to 3)')
add_bullet(doc, 'View Analytics Dashboard')
add_bullet(doc, 'Sign Up / Log In')

add_body(doc, 'Actor: Student (Authenticated)')
add_bullet(doc, 'All unauthenticated use cases')
add_bullet(doc, 'Save/Bookmark Universities')
add_bullet(doc, 'View Profile with Saved Universities')
add_bullet(doc, 'Log Out')

add_body(doc, 'Actor: System')
add_bullet(doc, 'Validate credentials and issue JWT tokens')
add_bullet(doc, 'Serve university data from MongoDB')
add_bullet(doc, 'Protect profile routes via JWT middleware')

add_body(doc, '[Use-case Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 33-34: 4.3.3 Sequence Diagram ══════════════════

page_break(doc)

add_heading_para(doc, '4.3.3 Sequence Diagram')

add_body(doc,
    'The Sequence diagrams map the chronological message flow between system components for key workflows.')

add_body(doc,
    'Authentication Sequence (Login): The User enters credentials on the Login page → React sends POST '
    '/api/auth/login to Express server → Server queries MongoDB for user by email → Server compares password '
    'hash using bcrypt → On success, server generates JWT token (24h expiry) → Token and user object returned '
    'to React → AuthContext stores token in localStorage and updates state → User is redirected to Home page.')

add_body(doc,
    'University Search Sequence: User enters exam type and rank on Home page → React navigates to Universities '
    'page with query parameters → UniversityContext fetches GET /api/universities from Express → Express queries '
    'MongoDB and returns sorted results → React applies client-side filters (rank eligibility, state, fees, '
    'NAAC grade) → Filtered results rendered as paginated UniversityCard grid.')

add_body(doc,
    'Save University Sequence: Authenticated user clicks bookmark icon on UniversityCard → SavedContext adds '
    'university ID to saved array → Array persisted to localStorage → Profile page reads SavedContext to display '
    'saved universities.')

add_body(doc, '[Sequence Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 35: 4.3.4 Activity Diagram ═════════════════════

page_break(doc)

add_heading_para(doc, '4.3.4 Activity Diagram')

add_body(doc,
    'Activity diagrams detail the operational workflows within the platform. The primary activity flow for '
    'university discovery is as follows:')

add_body(doc,
    'Start → User opens Home page → Selects Exam Type from dropdown → If "State CET" is selected, a nested '
    'sub-filter appears for selecting specific state → User enters Rank → Clicks "Find Universities" → System '
    'navigates to Universities page → UniversityContext loads data → Client-side filtering applied → Results '
    'displayed in paginated grid → User can View Details, Compare, or Save → End.')

add_body(doc,
    'Authentication Activity: Start → User clicks Login/Sign Up → Enters credentials → System validates → '
    'If valid: JWT issued, session created, redirect to Home → If invalid: error message displayed, retry → End.')

add_body(doc, '[Activity Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 36-37: 4.3.5 State Chart Diagram ═══════════════

add_heading_para(doc, '4.3.5 State Chart Diagram')

add_body(doc,
    'The State chart diagram visualizes the conditional states of the application and transitions between them.')

add_body(doc, 'Application States:')
add_bullet(doc, 'Idle: Application loaded, no user interaction.')
add_bullet(doc, 'LoggedOut: No JWT token present. Only public routes accessible.')
add_bullet(doc, 'Authenticating: User submitting login/signup form. Awaiting server response.')
add_bullet(doc, 'LoggedIn: Valid JWT in localStorage. AuthContext populated with user data. Profile and save features enabled.')
add_bullet(doc, 'FetchingData: UniversityContext calling GET /api/universities. Loading state displayed.')
add_bullet(doc, 'DataLoaded: University data cached in context. Filters, sorting, and pagination active.')
add_bullet(doc, 'ViewingDetails: User viewing a specific university detail page.')
add_bullet(doc, 'Comparing: User has added universities to comparison table.')
add_bullet(doc, 'Error: API call failed or invalid token. Error messages displayed. Fallback to LoggedOut if authentication error.')

add_body(doc, '[State Chart Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 38-39: 4.3.6 Object Diagram ════════════════════

page_break(doc)

add_heading_para(doc, '4.3.6 Object Diagram')

add_body(doc,
    'Object diagrams show instances of classes at a specific runtime moment, illustrating the actual data '
    'flowing through the system.')

add_body(doc,
    'Example University Object Instance (IIT Bombay): name = "Indian Institute of Technology Bombay", '
    'location = "Mumbai, Maharashtra", ranking = 1, fees = "2.5L", average_placement = "25 LPA", '
    'placement_percentage = "95%", type = "Public", required_exam = "JEE Advanced", '
    'average_rank_required = "Top 500", branches = [{name: "Computer Science", cutoffRank: "68"}, '
    '{name: "Electrical Engineering", cutoffRank: "341"}, ...].')

add_body(doc,
    'Example User Object Instance: name = "Bharath Kumar", email = "bharath@example.com", '
    'password = "$2a$10$...(hashed)", joined = "Apr 2026", reviews = 0.')

add_body(doc,
    'Example SavedContext State: savedIds = ["681a2f3e...", "681b4c7a..."] (array of MongoDB ObjectIds '
    'stored in localStorage).')

add_body(doc, '[Object Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 40: 4.3.7 Collaboration Diagram ════════════════

add_heading_para(doc, '4.3.7 Collaboration Diagram')

add_body(doc,
    'Collaboration diagrams emphasize the structural organization of objects that send and receive messages, '
    'reinforcing the interactions shown in the sequence diagrams but focusing on system topology.')

add_body(doc,
    'Key collaboration paths: (1) User → Navbar → SearchBar → Universities Page (search query flow), '
    '(2) Home Page → UniversityContext → Express Server → MongoDB (data fetch flow), '
    '(3) UniversityCard → SavedContext → localStorage (bookmark flow), '
    '(4) Login Page → AuthContext → Express Server → MongoDB → JWT → localStorage (auth flow).')

add_body(doc, '[Collaboration Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 41: 4.3.8 Component Diagram ════════════════════

page_break(doc)

add_heading_para(doc, '4.3.8 Component Diagram')

add_body(doc,
    'Component diagrams illustrate the physical modules of the system and their dependencies.')

add_body(doc, 'Frontend Components:')
add_bullet(doc, 'React Application Bundle (pages, components, context providers, CSS modules)')
add_bullet(doc, 'Recharts Library (data visualization — Bar, Pie, Scatter charts)')
add_bullet(doc, 'React Router DOM (client-side routing with protected routes)')

add_body(doc, 'Backend Components:')
add_bullet(doc, 'Express Web Server (REST API endpoints, authentication middleware)')
add_bullet(doc, 'Mongoose ODM (schema definitions, database queries)')
add_bullet(doc, 'bcryptjs (password hashing) and jsonwebtoken (JWT generation/verification)')

add_body(doc, 'Data Components:')
add_bullet(doc, 'MongoDB Database Server (universities and users collections)')
add_bullet(doc, 'JSON Dataset (universities.json — master data source)')
add_bullet(doc, 'Excel Spreadsheets (branch cutoff source data)')

add_body(doc, '[Component Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 42-43: 4.3.9 Deployment Diagram ════════════════

add_heading_para(doc, '4.3.9 Deployment Diagram')

add_body(doc,
    'The Deployment diagram represents the physical deployment architecture of the UniSearch India platform.')

add_body(doc, 'Client Tier:')
add_bullet(doc, 'User\'s web browser (Chrome, Edge, Firefox) runs the React SPA built by Create React App.')

add_body(doc, 'Static Hosting (Frontend):')
add_bullet(doc, 'GitHub Pages or Vercel serves the production build/ folder containing HTML, CSS, JS bundles.')

add_body(doc, 'Application Server (Backend):')
add_bullet(doc, 'Render or Railway hosts the Node.js + Express server on port 5000, handling REST API requests and JWT authentication.')

add_body(doc, 'Database Server:')
add_bullet(doc, 'MongoDB Atlas (cloud-hosted) stores the universities and users collections with automatic backups and scaling.')

add_body(doc,
    'Communication: Browser ↔ Static Host (HTTPS), Browser ↔ API Server (HTTP/HTTPS fetch), API Server ↔ '
    'MongoDB Atlas (Mongoose connection string). The decoupled architecture allows independent scaling and '
    'maintenance of each tier.')

add_body(doc, '[Deployment Diagram Figure — refer to attached diagrams]')


# ═══════════════════════ PAGE 44: 5. Implementation and Testing ══════════════

page_break(doc)

add_heading_para(doc, '5. Implementation and Testing')

add_body(doc,
    'This chapter covers the implementation details of the platform codebase and the various testing strategies '
    'employed to ensure reliability, performance, and security across both the API and user interface layers.')


# ═══════════════════════ PAGE 44-46: 5.1 Implementation of Codes ═════════════

add_heading_para(doc, '5.1 Implementation of Codes')

add_body(doc,
    'The codebase follows a modular pattern. The Node.js backend employs MVC principles separating routes, '
    'controllers, and models. The React frontend follows a component-based directory structure grouping by '
    'pages, services, and context providers.')

add_body(doc, 'Backend Implementation (server.js):')
add_body(doc,
    'The Express server initializes with CORS middleware and JSON body parsing. It connects to MongoDB using '
    'Mongoose with the connection string from environment variables (defaulting to localhost:27017). The '
    'authentication middleware extracts JWT from the Authorization header, verifies it, and attaches the decoded '
    'user to the request object. Five API routes are defined: POST /api/auth/signup (validates input, checks '
    'for duplicate email, hashes password with bcrypt, saves user, returns JWT), POST /api/auth/login (finds '
    'user by email, compares password, returns JWT), GET /api/users/profile (protected, returns user without '
    'password), GET /api/universities (returns all sorted by ranking), GET /api/universities/:id (returns single '
    'university by ObjectId).')

add_body(doc, 'Frontend Implementation:')
add_body(doc,
    'App.js serves as the root component wrapping all pages in AuthProvider, UniversityProvider, and '
    'SavedProvider context providers. React Router DOM v6 manages routing with protected routes redirecting '
    'unauthenticated users. The Universities page implements complex client-side filtering logic supporting '
    'multi-select checkboxes for exam type, state, institution type, NAAC grade, and fee range, along with '
    'rank-based eligibility checking. Pagination divides results into groups of 9 with prev/next navigation.')

add_body(doc, 'Database Seeding (seed.js):')
add_body(doc,
    'The seeding script reads from src/data/universities.json, clears the existing collection, enriches records '
    'with computed fallback values (fees estimated by institution type, placements estimated by ranking tier, '
    'required_exam derived from ranking and type), generates 7 default engineering branches with cutoffs scaled '
    'by ranking, and inserts all enriched records into MongoDB.')


# ═══════════════════════ PAGE 47: 5.2 Functional Testing ═════════════════════

page_break(doc)

add_heading_para(doc, '5.2 Functional Testing')

add_body(doc,
    'Functional tests validate that specific actions produce correct results. The following test cases '
    'were verified:')

func_tests = [
    'Rank-Based Search: Entering "Top 500" in JEE Advanced correctly returns only premier IITs within that cutoff bracket.',
    'Multi-Filter Combination: Selecting "Public" type + "Maharashtra" state + "Under 1L" fees returns only matching universities.',
    'Pagination: With 50+ results, page navigation correctly shows 9 cards per page with appropriate prev/next buttons.',
    'University Details: Clicking "View Details" on a card navigates to the correct university page with all metrics and branch cutoffs displayed.',
    'Comparison: Adding 3 universities to compare displays all metrics side-by-side; adding a 4th is restricted.',
    'Authentication: Signing up with valid credentials creates an account and returns a valid JWT; attempting duplicate email returns error.',
    'Bookmarking: Saving a university adds it to the profile page; unsaving removes it. Data persists across page reloads via localStorage.',
    'Global Search: Typing a university name in the Navbar search bar and pressing Enter navigates to the Universities page with pre-filtered results.',
]
for ft in func_tests:
    add_bullet(doc, ft)


# ═══════════════════════ PAGE 48: 5.3 Non-functional Testing ═════════════════

page_break(doc)

add_heading_para(doc, '5.3 Non-functional Testing')

add_body(doc,
    'Non-functional testing measures system performance and usability beyond core features:')

nf_tests = [
    'Page Load Time: Initial page load completes within 2-3 seconds on a standard broadband connection.',
    'API Response Time: GET /api/universities returns all records within 500ms from MongoDB.',
    'UI Rendering: Loading 9 paginated university cards renders completely within 200ms after data arrival.',
    'Responsiveness: The interface adapts correctly across desktop (1920px), tablet (768px), and mobile (375px) viewports.',
    'Concurrent Users: The Express server handles multiple simultaneous requests without degradation.',
    'Context Caching: UniversityContext caches the university list after the first fetch, avoiding redundant API calls on navigation between pages.',
]
for nft in nf_tests:
    add_bullet(doc, nft)


# ═══════════════════════ PAGE 49: 5.4 Error Handling Testing ═════════════════

add_heading_para(doc, '5.4 Error Handling Testing')

add_body(doc, 'Error handling test scenarios cover graceful degradation under failure conditions:')

err_tests = [
    'Invalid Login: Entering wrong email/password returns "Invalid credentials" message without exposing system details.',
    'Expired Token: Accessing /api/users/profile with an expired JWT returns 403 Forbidden and redirects to login.',
    'Invalid University ID: Navigating to /university/invalid-id returns a 404 "University not found" message.',
    'Network Failure: When the API server is unreachable, the frontend displays user-friendly error messages.',
    'Missing Required Fields: Signup with empty name/email/password returns 400 "All fields are required".',
    'Duplicate Email: Attempting signup with an existing email returns "User already exists with this email".',
]
for et in err_tests:
    add_bullet(doc, et)


# ═══════════════════════ PAGE 49: 5.5 Integration Testing ════════════════════

page_break(doc)

add_heading_para(doc, '5.5 Integration Testing')

add_body(doc,
    'Integration testing validates the data flow between system layers — from MongoDB through the Express '
    'controller to the React UI component.')

int_tests = [
    'Database to API: Seeded university data in MongoDB is correctly returned by GET /api/universities with proper JSON structure.',
    'API to Frontend: UniversityContext successfully fetches, parses, and stores the API response for rendering.',
    'Auth Flow End-to-End: Signup → Login → Profile access → Logout cycle works seamlessly with token persistence across reloads.',
    'Filter Pipeline: Query parameters from the Home page search form are correctly passed to the Universities page and applied as pre-selected filters.',
    'Save Sync: Universities saved via UniversityCard\'s bookmark toggle are correctly reflected on the Profile page through SavedContext.',
]
for it in int_tests:
    add_bullet(doc, it)


# ═══════════════════════ PAGE 50-52: 5.6 Unit Testing ════════════════════════

add_heading_para(doc, '5.6 Unit Testing')

add_body(doc,
    'Unit tests target individual functions and components in isolation to verify correctness of business logic:')

unit_tests = [
    'Password Hashing: bcrypt.hash() produces a valid hash; bcrypt.compare() correctly validates matching and non-matching passwords.',
    'JWT Token: jwt.sign() generates a token with correct payload (id, email); jwt.verify() successfully decodes valid tokens and rejects tampered ones.',
    'Authentication Middleware: authenticateToken correctly extracts Bearer token from Authorization header, rejects requests without tokens (401), and rejects expired/invalid tokens (403).',
    'Mongoose Models: University model validates required fields (name, location, ranking) and rejects documents missing them. User model enforces unique email constraint.',
    'Seeding Logic: Fallback estimation correctly assigns fees ("2.5L" for top public, "15L" for private), placement percentages based on ranking tier, and required_exam based on ranking and type.',
    'React Component Rendering: UniversityCard renders correctly with provided props (name, location, fees, ranking). Bookmark toggle switches between saved and unsaved states.',
    'Filter Logic: The filtering function in Universities.js correctly intersects multiple filter criteria (exam type AND state AND fee range) and returns only matching universities.',
    'Pagination Logic: Given 27 results with pageSize=9, the paginator correctly produces 3 pages and displays correct items for each page index.',
]
for ut in unit_tests:
    add_bullet(doc, ut)


# ═══════════════════════ PAGE 53-59: 6. Data Visualization ═══════════════════

page_break(doc)

add_heading_para(doc, '6. Data Visualization')

add_body(doc,
    'The UniSearch India platform incorporates a comprehensive analytics dashboard that provides visual insights '
    'into the university dataset using Recharts, a composable charting library built on React components. '
    'The Dashboard page is accessible to all users and presents four key visualizations.')

add_body(doc, 'Summary Statistics Cards:')
add_body(doc,
    'The dashboard begins with three summary cards displaying: Total Ranked Universities (count of all universities '
    'in the dataset), States Covered (number of unique states represented), and Average Top 10 Placement (mean '
    'placement percentage of the top 10 NIRF-ranked universities). These cards provide an at-a-glance overview '
    'of the dataset scope.')

add_body(doc, 'Bar Chart — Universities per State (Top 10):')
add_body(doc,
    'A vertical bar chart displays the top 10 states by number of ranked universities. The x-axis represents '
    'states and the y-axis represents the count of universities. This visualization helps students identify states '
    'with the highest concentration of ranked institutions, which is useful for location-based decision-making. '
    'The chart dynamically computes state distribution from the dataset using array aggregation.')

add_body(doc, 'Pie Chart (Donut) — NAAC Accreditation Distribution:')
add_body(doc,
    'A donut-style pie chart shows the distribution of universities by NAAC accreditation grade (A++, A+, A). '
    'Each segment is color-coded and labeled with the percentage. This chart helps students quickly assess the '
    'quality landscape — for instance, how many universities hold the highest A++ accreditation versus A or A+. '
    'The data is computed by grouping universities by their NAAC grade field.')

add_body(doc, 'Scatter Chart — Fees vs. Placement Probability:')
add_body(doc,
    'A scatter plot visualizes the relationship between annual tuition fees (x-axis) and placement percentage '
    '(y-axis). Each dot represents a university. This visualization reveals an important insight: higher fees '
    'do not always correlate with better placement outcomes. Many public universities with lower fees achieve '
    'comparable or superior placement rates to expensive private institutions. This data-driven insight helps '
    'students make cost-effective decisions.')

add_body(doc, 'Call-to-Action Section:')
add_body(doc,
    'Below the charts, the dashboard includes CTA buttons linking to the Compare page and the University Search '
    'Database, encouraging users to take action based on the analytical insights they have gained.')

add_body(doc, 'Technical Implementation:')
add_body(doc,
    'All charts are rendered using Recharts components (BarChart, PieChart, ScatterChart) with responsive '
    'containers. The data is derived from the UniversityContext, which caches the full university list. '
    'Aggregation logic (state grouping, NAAC distribution, fee-placement mapping) is computed client-side '
    'using JavaScript array methods (reduce, filter, map). The dashboard updates automatically when the '
    'underlying dataset changes.')

add_body(doc, '[Dashboard Screenshots — refer to attached figures]')


# ═══════════════════════ PAGE 60-62: 7. Result ═══════════════════════════════

page_break(doc)

add_heading_para(doc, '7. Result')

add_body(doc,
    'The University Recommendation and Comparison Platform has been successfully developed and tested as a '
    'fully functional full-stack web application. The following results were achieved:')

results = [
    'Rank-Based Discovery: Students can input their entrance exam type and rank on the Home page and instantly receive a filtered list of eligible universities. The system accurately matches ranks against branch-wise cutoff data for JEE Main, JEE Advanced, CUET, State CET, and Institute Specific exams.',
    'Advanced Filtering: The universities listing page supports multi-criteria filtering by exam type, state, institution type (Public/Private), NAAC grade, and fee range. Filters work in combination and update results in real-time with pagination.',
    'Detailed Profiles: Each university page displays comprehensive metrics including NIRF ranking, annual fees, average placement packages, placement percentages, branch-wise cutoff tables, and links to official websites.',
    'Side-by-Side Comparison: The comparison feature allows students to compare up to 3 universities simultaneously across 9 metrics (location, NAAC grade, NIRF ranking, fees, placement %, average package, established year, required exam, average rank required).',
    'Analytics Dashboard: Three interactive charts (bar, donut pie, scatter) provide visual insights into state distribution, accreditation quality, and the fee-placement relationship.',
    'Secure Authentication: JWT-based signup/login with bcrypt password hashing protects user accounts. Sessions persist across page reloads. Protected routes correctly restrict access.',
    'Bookmarking and Profile: Authenticated users can save universities and manage their shortlist from the Profile page. Data persists via localStorage.',
    'Performance: API responses average under 500ms. Page transitions are smooth with Context-based caching eliminating redundant API calls.',
]
for r in results:
    add_bullet(doc, r)

add_body(doc,
    'The platform has been deployed with the frontend on GitHub Pages and is architecturally ready for '
    'full-stack deployment. All functional and non-functional requirements specified in Chapter 3 have been '
    'met successfully.')


# ═══════════════════════ PAGE 60-62: 8. Conclusion ═══════════════════════════

page_break(doc)

add_heading_para(doc, '8. Conclusion and Future Enhancement')

add_body(doc,
    'The Integrated Information Platform for Indian Universities (UniSearch India) is a meaningful and '
    'technically rich solution to a real-world problem faced by students during the admission process. By '
    'combining structured data, intelligent filtering, secure authentication, and interactive analytics, '
    'the platform empowers users to make informed and confident decisions.')


# ═══════════════════════ PAGE 63: 8.1 Conclusion ═════════════════════════════

add_heading_para(doc, '8.1 Conclusion')

add_body(doc,
    'The project demonstrates that aggregating scattered educational data into a singular, interconnected '
    'platform provides immense utility to aspiring engineering students. The platform successfully bridges '
    'the information gap by providing rank-based filtering, multi-criteria search, side-by-side comparison, '
    'and visual analytics — all within a secure and user-friendly interface.')

add_body(doc,
    'From a technical standpoint, the project showcases the integration of modern web technologies — React '
    'for the frontend, Node.js/Express for the backend, MongoDB for data persistence, and JWT/bcrypt for '
    'security. The modular architecture ensures maintainability, and the decoupled deployment model supports '
    'scalability.')

add_body(doc,
    'This project is not only a strong academic submission but also a valuable product concept with real '
    'deployment potential. It showcases full-stack development skills, thoughtful system design, and a clear '
    'understanding of user needs — making it an excellent example of technology solving a socially relevant '
    'challenge.')


# ═══════════════════════ PAGE 63-64: 8.2 Future Enhancement ═════════════════

add_heading_para(doc, '8.2 Future Enhancement')

add_body(doc,
    'Although the current version is fully functional, the platform has significant scope for future '
    'improvements:')

future = [
    'Integration of machine learning models for personalized recommendations based on student profiles and historical admission patterns.',
    'Real-time cutoff prediction based on historical trends, counseling round data, and seat availability.',
    'Category-wise and quota-based admission logic to support reservation categories (SC/ST/OBC/EWS).',
    'Push notifications for counseling round updates and important admission deadlines.',
    'Admin dashboard for dynamic data management — allowing administrators to add, update, and delete university records without code changes.',
    'Export options including PDF reports and shortlist summaries that students can share or print.',
    'Mobile application version (React Native) for wider accessibility across Android and iOS devices.',
    'Integration with university application portals for direct application submission from the platform.',
]
for f in future:
    add_bullet(doc, f)

add_body(doc,
    'These enhancements can transform the project from an academic submission into a complete education '
    'guidance ecosystem serving thousands of students across India.')


# ═══════════════════════ PAGE 64+: 9. References ═════════════════════════════

page_break(doc)

add_heading_para(doc, '9. References')

refs = [
    '[1] React Documentation — https://react.dev/ — Official documentation for React 18, the UI library used for building the single-page application frontend.',
    '[2] Node.js Documentation — https://nodejs.org/en/docs — Official reference for the Node.js runtime environment powering the Express backend server.',
    '[3] Express.js Guide — https://expressjs.com/ — Web framework for Node.js used to build the RESTful API endpoints.',
    '[4] MongoDB Manual — https://www.mongodb.com/docs/manual/ — Official documentation for MongoDB, the NoSQL database used for persistent data storage.',
    '[5] Mongoose ODM — https://mongoosejs.com/docs/ — Object Data Modeling library for MongoDB, used for schema definitions and database queries.',
    '[6] JSON Web Tokens (JWT) — https://jwt.io/introduction — Reference for JWT-based stateless authentication implemented in the platform.',
    '[7] bcrypt.js — https://github.com/dcodeIO/bcrypt.js — Library used for secure password hashing using the bcrypt algorithm.',
    '[8] Recharts — https://recharts.org/ — Composable charting library for React used in the analytics dashboard (Bar, Pie, Scatter charts).',
    '[9] React Router DOM v6 — https://reactrouter.com/ — Client-side routing library used for navigation and protected route implementation.',
    '[10] NIRF Rankings, Ministry of Education, Government of India — https://www.nirfindia.org/ — Source of university ranking data used in the platform dataset.',
    '[11] JoSAA (Joint Seat Allocation Authority) — https://josaa.nic.in/ — Reference for JEE Main/Advanced cutoff rank data and admission processes.',
    '[12] MongoDB Atlas — https://www.mongodb.com/atlas — Cloud-hosted database service used for production deployment.',
    '[13] GitHub Pages — https://pages.github.com/ — Static hosting platform used for frontend deployment.',
    '[14] Create React App — https://create-react-app.dev/ — Project scaffolding and build tooling used for the React application.',
]
for ref in refs:
    add_body(doc, ref)


# ── save ─────────────────────────────────────────────────────────────────────

output_path = 'Project_Report.docx'
doc.save(output_path)
print(f'Successfully generated {output_path}')
